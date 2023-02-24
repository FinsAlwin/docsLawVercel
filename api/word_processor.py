from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor


class WordProcessor:
    def __init__(self):
        self

    def getDocx(self, path):
        if path:
            doc = Document(path)
            return doc
        else:
            doc = Document()
            return doc

    def addPageBreak(self, doc):
        last_element = doc.element.body[-1]
        if last_element.tag.endswith('p'):
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.add_run().add_break(WD_BREAK.PAGE)

    def addToDocx(self, doc, dest_doc):
        for element in doc.element.body:
            if element.tag.endswith(('p', 'tbl')):
                new_element = element
                dest_doc.element.body.append(new_element)

    def saveTolocal(self, dest_doc, name):
        dest_doc.save(name)


class DocxEditor:
    def __init__(self):
        self

    def replace_placeholders(self, document, variables):
        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                for placeholder, value in variables.items():
                    if placeholder in paragraph.text:
                        run.text = run.text.replace(
                            placeholder, f'{value.upper()}'
                        )

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            for key, value in variables.items():
                                if key in run.text:
                                    run.text = run.text.replace(
                                        key, value).upper()

    def add_index_table_row(self, table: Table, cells_text: list) -> None:

        new_row = table.add_row()
        for i, cell_text in enumerate(cells_text):
            new_row.cells[i].text = f'  {cell_text.upper()}'

            if i == 0:
                for paragraph in new_row.cells[i].paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Format the cell with a border
            tc = new_row.cells[i]._element
            tcPr = OxmlElement('w:tcPr')
            tc.append(tcPr)
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)
            top = OxmlElement('w:top')
            top.set(qn('w:val'), "single")
            top.set(qn('w:sz'), "6")
            top.set(qn('w:space'), "0")
            top.set(qn('w:color'), "000000")
            tcBorders.append(top)
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), "single")
            bottom.set(qn('w:sz'), "6")
            bottom.set(qn('w:space'), "0")
            bottom.set(qn('w:color'), "000000")
            tcBorders.append(bottom)
            left = OxmlElement('w:left')
            left.set(qn('w:val'), "single")
            left.set(qn('w:sz'), "6")
            left.set(qn('w:space'), "0")
            left.set(qn('w:color'), "000000")
            tcBorders.append(left)
            right = OxmlElement('w:right')
            right.set(qn('w:val'), "single")
            right.set(qn('w:sz'), "6")
            right.set(qn('w:space'), "0")
            right.set(qn('w:color'), "000000")
            tcBorders.append(right)

            # Add font, font size, and font style to the cell text
            font = new_row.cells[i].paragraphs[0].runs[0].font
            font.bold = False  # make the font bold
            font.name = "Bookman Old Style"  # set the font name to Arial
            font.size = Pt(13)  # set the font size to 12pt
            # set the font color to black
            font.color.rgb = RGBColor(0x00, 0x00, 0x00)
